# app.py
# ---------------------------------------------
# Personal Trainer x AI 診断フォーム（単一ファイル版）
# 条件対応：
# ① 分析結果を詳細&わかりやすく表示（BMR/TDEE/目的別カロリー・PFC、運動提案）
# ② フォーム情報はSQLiteに保存（/data.db）、CSV/Excel書き出し可能
# ③ AI姿勢チェック：写真アップロード→角度/左右差を算出して注意喚起（MediaPipe使用）
#    ※MediaPipe未インストール/失敗時は機能を自動で無効化
#    食事提案：目的/好み/アレルギーに合わせて1日の具体例を生成
# ④ 見込み客→予約のLINE配信テンプレをアプリ内で自動生成&DL
# ⑤ GitHub/Streamlit CloudにそのままデプロイOK（requirements.txt同梱）
# ---------------------------------------------

import os
import io
import math
import json
import sqlite3
from datetime import datetime
from typing import Dict, Any, Tuple

import numpy as np
import pandas as pd
from PIL import Image

import streamlit as st

# 姿勢チェックは MediaPipe があれば有効化
POSE_AVAILABLE = True
try:
    import mediapipe as mp
except Exception:
    POSE_AVAILABLE = False

# Wordレポート作成（日本語OK）
from docx import Document

APP_TITLE = "AIフィットネス診断 & 姿勢チェック（個人トレーナー向け）"
DB_PATH = "data.db"


# ============== DB ユーティリティ ==============
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
    CREATE TABLE IF NOT EXISTS leads (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        phone TEXT,
        age INTEGER,
        gender TEXT,
        height_cm REAL,
        weight_kg REAL,
        activity_level TEXT,
        goal TEXT,
        dietary_prefs TEXT,
        allergies TEXT,
        created_at TEXT
    )
    """)
    c.execute("""
    CREATE TABLE IF NOT EXISTS assessments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        lead_id INTEGER,
        bmi REAL,
        bmr REAL,
        tdee REAL,
        target_calories REAL,
        protein_g REAL,
        fat_g REAL,
        carbs_g REAL,
        notes TEXT,
        posture_findings TEXT,
        created_at TEXT,
        FOREIGN KEY (lead_id) REFERENCES leads(id)
    )
    """)
    conn.commit()
    conn.close()


def insert_lead_and_assessment(lead: Dict[str, Any], assess: Dict[str, Any]) -> int:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""
    INSERT INTO leads (name, email, phone, age, gender, height_cm, weight_kg, activity_level, goal, dietary_prefs, allergies, created_at)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        lead["name"], lead["email"], lead["phone"], lead["age"], lead["gender"],
        lead["height_cm"], lead["weight_kg"], lead["activity_level"], lead["goal"],
        lead["dietary_prefs"], lead["allergies"], lead["created_at"]
    ))
    lead_id = c.lastrowid
    c.execute("""
    INSERT INTO assessments (lead_id, bmi, bmr, tdee, target_calories, protein_g, fat_g, carbs_g, notes, posture_findings, created_at)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        lead_id, assess["bmi"], assess["bmr"], assess["tdee"], assess["target_calories"],
        assess["protein_g"], assess["fat_g"], assess["carbs_g"], assess["notes"],
        assess.get("posture_findings", ""), assess["created_at"]
    ))
    conn.commit()
    conn.close()
    return lead_id


def load_all_data():
    conn = sqlite3.connect(DB_PATH)
    leads = pd.read_sql_query("SELECT * FROM leads ORDER BY created_at DESC", conn)
    assessments = pd.read_sql_query("SELECT * FROM assessments ORDER BY created_at DESC", conn)
    conn.close()
    return leads, assessments


# ============== 栄養計算ユーティリティ ==============
def calc_bmi(weight_kg: float, height_cm: float) -> float:
    h_m = height_cm / 100.0
    return round(weight_kg / (h_m ** 2), 2)


def mifflin_st_jeor_bmr(gender: str, weight_kg: float, height_cm: float, age: int) -> float:
    # 男性: 10W + 6.25H - 5A + 5, 女性: 10W + 6.25H - 5A - 161
    if gender == "男性":
        return 10 * weight_kg + 6.25 * height_cm - 5 * age + 5
    else:
        return 10 * weight_kg + 6.25 * height_cm - 5 * age - 161


def activity_factor(level: str) -> float:
    mapping = {
        "低い（デスクワーク中心/運動ほぼ無し）": 1.2,
        "やや低い（週1〜2軽い運動）": 1.375,
        "普通（週3〜4運動）": 1.55,
        "高い（週5以上ハード）": 1.725,
        "非常に高い（アスリート級）": 1.9,
    }
    return mapping.get(level, 1.2)


def target_calories_from_goal(tdee: float, goal: str) -> float:
    if goal == "減量（-15〜20%）":
        return round(tdee * 0.85)
    elif goal == "緩やか減量（-10%）":
        return round(tdee * 0.90)
    elif goal == "現状維持":
        return round(tdee)
    elif goal == "増量（+10%）":
        return round(tdee * 1.10)
    else:
        return round(tdee)


def macro_plan(weight_kg: float, calories: float, goal: str):
    # たんぱく質：1.8 g/kg（中庸）
    protein_g = 1.8 * weight_kg
    # 脂質：25%
    fat_kcal = calories * 0.25
    fat_g = fat_kcal / 9.0
    # 炭水化物：残り
    protein_kcal = protein_g * 4.0
    carbs_kcal = max(0.0, calories - (protein_kcal + fat_kcal))
    carbs_g = carbs_kcal / 4.0
    return round(protein_g), round(fat_g), round(carbs_g)


# ============== 食事提案（1日例） ==============
def meal_suggestions(cal: int, p: int, f: int, c: int, prefs: str, allergies: str, goal: str):
    avoid = [a.strip() for a in allergies.split(",") if a.strip()]
    lowfat = goal.startswith("減量")

    def ok(item: str) -> bool:
        return all(a.lower() not in item.lower() for a in avoid)

    breakfast = [i for i in [
        "オートミール+無糖ヨーグルト+ベリー",
        "全卵1+卵白2のスクランブル+玄米おにぎり",
        "プロテインシェイク+バナナ",
    ] if ok(i)]
    lunch = [i for i in [
        "鶏むねグリル150g+雑穀米150g+サラダ",
        "鮭の塩焼き+さつまいも200g+味噌汁",
        "豆腐ステーキ+玄米150g+野菜炒め",
    ] if ok(i)]
    dinner = [i for i in [
        "白身魚のホイル焼き+ブロッコリー+じゃがいも150g",
        "豚ヒレ100g+白菜スープ+玄米120g",
        "鶏つくね鍋（春雨少量）",
    ] if ok(i)]
    snack = [i for i in [
        "プロテインバー/和風おにぎり/枝豆/ミックスナッツ少量",
    ] if ok(i)]

    guide = f"""
- 1日の目標：{cal} kcal / P{p}g F{f}g C{c}g
- 配分例：朝 25% / 昼 35% / 夜 30% / 間食 10%
- 水分：目安 体重(kg)×30〜35 ml
- {"減量中：脂質控えめ・高たんぱくを意識" if lowfat else "増量/維持：炭水化物を運動前後に寄せる"} 
- 好み：{prefs or "（未入力）"} / アレルギー回避：{', '.join(avoid) if avoid else 'なし'}
""".strip()
    return {
        "breakfast": breakfast[:2] or ["（該当食材を回避して選択）"],
        "lunch": lunch[:2] or ["（該当食材を回避して選択）"],
        "dinner": dinner[:2] or ["（該当食材を回避して選択）"],
        "snack": snack[:1],
        "guide": guide
    }


# ============== 姿勢チェック（MediaPipe） ==============
def analyze_posture(image: Image.Image):
    """
    正面立位想定。
    - 肩ラインの傾き（左右差）：|angle| >= 5°
    - 骨盤（左右の腰）：|angle| >= 5°
    - 頭部の傾き（左右の耳）：|angle| >= 5°
    - ニーイン傾向：両膝間距離 / 両足首間距離 < 0.9 なら注意
    """
    findings = []
    if not POSE_AVAILABLE:
        return {"ok": False, "message": "MediaPipeが利用できないため、姿勢解析をスキップしました。", "findings": findings}

    mp_pose = mp.solutions.pose
    img = image.convert("RGB")
    arr = np.array(img)

    with mp_pose.Pose(static_image_mode=True) as pose:
        res = pose.process(arr)
        if not res.pose_landmarks:
            return {"ok": False, "message": "ランドマークを検出できませんでした。正面から全身〜上半身が写る明るい写真でお試しください。", "findings": findings}

        lm = res.pose_landmarks.landmark

        def get_xy(idx):
            return lm[idx].x, lm[idx].y

        def line_angle_deg(p1, p2):
            dx = p2[0] - p1[0]
            dy = p2[1] - p1[1]
            return math.degrees(math.atan2(dy, dx))  # 水平=0°

        def dist(p1, p2):
            return math.hypot(p2[0]-p1[0], p2[1]-p1[1])

        # 肩
        ls = get_xy(mp_pose.PoseLandmark.LEFT_SHOULDER.value)
        rs = get_xy(mp_pose.PoseLandmark.RIGHT_SHOULDER.value)
        shoulder_deg = line_angle_deg(rs, ls)  # 右→左
        # 骨盤
        lh = get_xy(mp_pose.PoseLandmark.LEFT_HIP.value)
        rh = get_xy(mp_pose.PoseLandmark.RIGHT_HIP.value)
        hip_deg = line_angle_deg(rh, lh)
        # 頭（耳ライン）
        le = get_xy(mp_pose.PoseLandmark.LEFT_EAR.value)
        re = get_xy(mp_pose.PoseLandmark.RIGHT_EAR.value)
        head_deg = line_angle_deg(re, le)
        # 膝と足首距離
        lk = get_xy(mp_pose.PoseLandmark.LEFT_KNEE.value)
        rk = get_xy(mp_pose.PoseLandmark.RIGHT_KNEE.value)
        la = get_xy(mp_pose.PoseLandmark.LEFT_ANKLE.value)
        ra = get_xy(mp_pose.PoseLandmark.RIGHT_ANKLE.value)
        knee_w = dist(lk, rk)
        ankle_w = dist(la, ra)

        if abs(shoulder_deg) >= 5:
            findings.append(f"肩の高さの左右差：{shoulder_deg:.1f}°（5°以上→要注意）")
        if abs(hip_deg) >= 5:
            findings.append(f"骨盤の左右差：{hip_deg:.1f}°（5°以上→要注意）")
        if abs(head_deg) >= 5:
            findings.append(f"頭部の傾き：{head_deg:.1f}°（5°以上→要注意）")
        if ankle_w > 0 and knee_w / ankle_w < 0.9:
            ratio = knee_w / ankle_w
            findings.append(f"ニーイン傾向（膝間/足首間比）：{ratio:.2f}（<0.90で注意）")

        advice = []
        if any("肩" in f for f in findings):
            advice.append("肩の左右差→僧帽筋上部の過緊張/腹斜筋の弱さの可能性。サイドプランク/ショルダープレスのフォーム修正。")
        if any("骨盤" in f for f in findings):
            advice.append("骨盤の左右差→中臀筋/大臀筋の弱さ、股関節の可動域不足。クラムシェル/ヒップヒンジ練習。")
        if any("頭部" in f for f in findings):
            advice.append("頭部の傾き→胸鎖乳突筋/僧帽筋の左右差。胸椎伸展と頸部の軽いストレッチを習慣化。")
        if any("ニーイン" in f for f in findings):
            advice.append("ニーイン→股関節外旋筋/内転筋バランス。チューブで膝外押し意識のスクワット、グルートブリッジ。")

        return {"ok": True, "message": "解析完了", "findings": findings, "advice": advice}


# ============== レポート（DOCX） ==============
def build_report_docx(lead: Dict[str, Any], assess: Dict[str, Any], meals, posture) -> bytes:
    doc = Document()
    doc.add_heading('AIフィットネス診断レポート', level=1)

    p = doc.add_paragraph()
    p.add_run(f"お名前：{lead['name']}").bold = True
    doc.add_paragraph(f"作成日時：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading('基本情報', level=2)
    doc.add_paragraph(f"年齢：{lead['age']} / 性別：{lead['gender']}")
    doc.add_paragraph(f"身長：{lead['height_cm']} cm / 体重：{lead['weight_kg']} kg")
    doc.add_paragraph(f"活動レベル：{lead['activity_level']} / 目標：{lead['goal']}")
    doc.add_paragraph(f"好み：{lead['dietary_prefs']} / アレルギー：{lead['allergies']}")

    doc.add_heading('分析結果（栄養）', level=2)
    doc.add_paragraph(f"BMI：{assess['bmi']}")
    doc.add_paragraph(f"BMR（基礎代謝）：{round(assess['bmr'])} kcal")
    doc.add_paragraph(f"TDEE（推定消費）：{round(assess['tdee'])} kcal")
    doc.add_paragraph(f"目標カロリー：{assess['target_calories']} kcal")
    doc.add_paragraph(f"PFC目標：P{assess['protein_g']}g / F{assess['fat_g']}g / C{assess['carbs_g']}g")

    doc.add_heading('1日の食事例', level=2)
    doc.add_paragraph(meals["guide"])
    doc.add_paragraph(f"朝：{', '.join(meals['breakfast'])}")
    doc.add_paragraph(f"昼：{', '.join(meals['lunch'])}")
    doc.add_paragraph(f"夜：{', '.join(meals['dinner'])}")
    doc.add_paragraph(f"間食：{', '.join(meals['snack'])}")

    doc.add_heading('姿勢チェック', level=2)
    if posture.get("ok"):
        if posture.get("findings"):
            doc.add_paragraph("注意ポイント：")
            for f in posture["findings"]:
                doc.add_paragraph(f"・{f}")
        else:
            doc.add_paragraph("大きな偏りは検出されませんでした。")
        if posture.get("advice"):
            doc.add_paragraph("改善アドバイス：")
            for a in posture["advice"]:
                doc.add_paragraph(f"・{a}")
    else:
        doc.add_paragraph(posture.get("message", "解析を実行していません。"))

    doc.add_heading('トレーニングの目安', level=2)
    doc.add_paragraph("週3〜4回、全身（スクワット/ヒンジ/プレス/ロー/体幹）を基本。フォームを動画で毎回チェック。")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ============== LINE ステップ配信テンプレ ==============
def build_line_step_template(lead_name_placeholder="（お名前）") -> str:
    return f"""【LINEステップ配信テンプレ（見込み客→予約）】

Step0（登録直後）：
{lead_name_placeholder}さん、登録ありがとうございます！
目的達成まで「短時間×最短距離」で並走します。明日、無料AI診断の結果ダイジェストを送りますね。

Step1（翌日朝9時）：
AI診断ダイジェスト：
・体型指標：BMI / 推定基礎代謝 / 1日の目安カロリー
・PFCバランス：高たんぱく・脂質25%・残り炭水化物
・最初の1週間：姿勢リセット＋軽めの全身メニュー
→ 詳細レポートが欲しい方は「レポート」と返信

Step2（翌日朝9時）：
姿勢チェックの注意点まとめ：
・肩/骨盤の左右差、頭部の傾き、膝の入り込み など
・改善ドリル（1日5分）
→ ご希望なら無料15分のフォーム相談を実施。「相談」と返信

Step3（翌日朝9時）：
食事テンプレ：
・朝：オートミール＋無糖ヨーグルト＋ベリー
・昼：鶏むね＋雑穀米＋サラダ
・夜：白身魚ホイル焼き＋温野菜＋じゃがいも
→ 個別最適の食事プランは「食事」と返信

Step4（翌日朝9時）：
【先着5名】体験セッション（30分/オンライン）0円
・AI診断に基づく個別メニュー提案
・今の課題3つを可視化
ご希望の日時を第3候補まで返信ください。記入例：
「第1：10/20 20:00 第2：10/21 21:30 第3：10/22 19:00」
"""


# ============== Streamlit 画面 ==============
def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="💪", layout="wide")
    st.title(APP_TITLE)
    st.caption("※ 本アプリは学習/提案目的です。医療的助言が必要な場合は専門家へ相談してください。")

    init_db()

    tabs = st.tabs(["1) フォーム入力", "2) 姿勢チェック", "3) 結果&レポート", "4) 保存データ/書き出し", "5) LINEステップ"])

    # ==== 1) フォーム入力 ====
    with tabs[0]:
        st.subheader("お客様情報")
        with st.form("lead_form", clear_on_submit=False):
            c1, c2, c3 = st.columns(3)
            name = c1.text_input("お名前", "")
            email = c2.text_input("メールアドレス", "")
            phone = c3.text_input("電話番号（任意）", "")

            c4, c5, c6, c7 = st.columns(4)
            age = c4.number_input("年齢", 10, 100, 35)
            gender = c5.selectbox("性別", ["男性", "女性"])
            height_cm = c6.number_input("身長（cm）", 120, 220, 170)
            weight_kg = c7.number_input("体重（kg）", 30.0, 200.0, 65.0, step=0.1)

            activity_level = st.selectbox("活動レベル", [
                "低い（デスクワーク中心/運動ほぼ無し）",
                "やや低い（週1〜2軽い運動）",
                "普通（週3〜4運動）",
                "高い（週5以上ハード）",
                "非常に高い（アスリート級）",
            ])
            goal = st.selectbox("目標", ["減量（-15〜20%）", "緩やか減量（-10%）", "現状維持", "増量（+10%）"])
            dietary_prefs = st.text_input("食の好み（例：和食/高たんぱく/低脂質など）", "")
            allergies = st.text_input("アレルギー（カンマ区切り。例：乳, 卵, 小麦）", "")

            submitted = st.form_submit_button("診断を実行")
            if submitted:
                try:
                    bmi = calc_bmi(weight_kg, height_cm)
                    bmr = mifflin_st_jeor_bmr(gender, weight_kg, height_cm, age)
                    tdee_val = bmr * activity_factor(activity_level)
                    target_cal = target_calories_from_goal(tdee_val, goal)
                    p, f, c = macro_plan(weight_kg, target_cal, goal)
                    notes = "高たんぱく/野菜多め/水分を十分に。週3〜4の全身トレと十分な睡眠を推奨。"

                    st.session_state["lead"] = {
                        "name": name.strip() or "匿名",
                        "email": email.strip(),
                        "phone": phone.strip(),
                        "age": int(age),
                        "gender": gender,
                        "height_cm": float(height_cm),
                        "weight_kg": float(weight_kg),
                        "activity_level": activity_level,
                        "goal": goal,
                        "dietary_prefs": dietary_prefs.strip(),
                        "allergies": allergies.strip(),
                        "created_at": datetime.now().isoformat(timespec="seconds"),
                    }
                    st.session_state["assess"] = {
                        "bmi": bmi,
                        "bmr": bmr,
                        "tdee": tdee_val,
                        "target_calories": target_cal,
                        "protein_g": p,
                        "fat_g": f,
                        "carbs_g": c,
                        "notes": notes,
                        "created_at": datetime.now().isoformat(timespec="seconds"),
                    }
                    st.success("診断を実行しました。タブ『2) 姿勢チェック』or『3) 結果&レポート』へ。")

                except Exception as e:
                    st.error(f"診断に失敗しました：{e}")

    # ==== 2) 姿勢チェック ====
    with tabs[1]:
        st.subheader("AI姿勢チェック（正面写真を1枚）")
        if POSE_AVAILABLE:
            img_file = st.file_uploader("画像（.jpg/.png）をアップロード", type=["jpg", "jpeg", "png"])
            if img_file is not None:
                try:
                    image = Image.open(img_file)
                    st.image(image, caption="アップロード画像プレビュー", use_column_width=True)
                    if st.button("姿勢を解析する"):
                        result = analyze_posture(image)
                        st.session_state["posture"] = result
                        if result["ok"]:
                            st.success("姿勢解析：完了")
                            if result["findings"]:
                                st.write("**注意ポイント**")
                                for f in result["findings"]:
                                    st.write(f"- {f}")
                            if result.get("advice"):
                                st.write("**改善アドバイス**")
                                for a in result["advice"]:
                                    st.write(f"- {a}")
                        else:
                            st.warning(result.get("message", "解析に失敗しました。"))
                except Exception as e:
                    st.error(f"画像解析でエラーが発生しました：{e}")
        else:
            st.info("MediaPipeが未利用のため、姿勢解析は利用できません。requirements.txt 通りにデプロイすると有効化されます。")

    # ==== 3) 結果&レポート ====
    with tabs[2]:
        st.subheader("結果（栄養・食事・姿勢）とレポート作成")
        lead = st.session_state.get("lead")
        assess = st.session_state.get("assess")
        posture = st.session_state.get("posture", {"ok": False, "message": "未解析", "findings": []})

        if not lead or not assess:
            st.warning("まず『1) フォーム入力』で診断を実行してください。")
        else:
            meals = meal_suggestions(
                int(assess["target_calories"]),
                int(assess["protein_g"]),
                int(assess["fat_g"]),
                int(assess["carbs_g"]),
                lead["dietary_prefs"],
                lead["allergies"],
                lead["goal"]
            )

            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("BMI", assess["bmi"])
                st.metric("BMR（kcal）", round(assess["bmr"]))
            with c2:
                st.metric("TDEE（kcal）", round(assess["tdee"]))
                st.metric("目標カロリー", assess["target_calories"])
            with c3:
                st.metric("P（g）", assess["protein_g"])
                st.metric("F（g） / C（g）", f"{assess['fat_g']} / {assess['carbs_g']}")

            st.write("### 食事の指針（1日サンプル）")
            st.write(meals["guide"])
            st.write(f"- 朝：{', '.join(meals['breakfast'])}")
            st.write(f"- 昼：{', '.join(meals['lunch'])}")
            st.write(f"- 夜：{', '.join(meals['dinner'])}")
            st.write(f"- 間食：{', '.join(meals['snack'])}")

            st.write("### 姿勢チェック結果")
            if posture.get("ok"):
                if posture.get("findings"):
                    for f in posture["findings"]:
                        st.write(f"- {f}")
                else:
                    st.write("- 大きな偏りは検出されませんでした。")
                if posture.get("advice"):
                    st.write("**改善アドバイス**")
                    for a in posture["advice"]:
                        st.write(f"- {a}")
            else:
                st.write(f"- {posture.get('message', '未解析')}")

            st.write("---")
            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button("この結果をDBに保存する"):
                    try:
                        lead_id = insert_lead_and_assessment(lead, {
                            **assess,
                            "posture_findings": json.dumps(posture, ensure_ascii=False)
                        })
                        st.success(f"保存しました（Lead ID: {lead_id}）")
                    except Exception as e:
                        st.error(f"保存に失敗しました：{e}")

            with c2:
                try:
                    docx_bytes = build_report_docx(lead, assess, meals, posture)
                    st.download_button(
                        "顧客用レポート（DOCX）をダウンロード",
                        data=docx_bytes,
                        file_name=f"report_{lead['name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"レポート生成に失敗しました：{e}")

            with c3:
                st.info("※ PDFでなくDOCXにした理由：日本語フォントの同梱が不要で確実に文字化けを回避できるため。")

    # ==== 4) 保存データ/書き出し ====
    with tabs[3]:
        st.subheader("保存データの一覧とエクスポート")
        leads_df, assess_df = load_all_data()
        st.write("**Leads**")
        st.dataframe(leads_df, use_container_width=True)
        st.write("**Assessments**")
        st.dataframe(assess_df, use_container_width=True)

        # Export CSV
        c1, c2 = st.columns(2)
        with c1:
            csv_bytes = leads_df.to_csv(index=False).encode("utf-8-sig")
            st.download_button("LeadsをCSVでダウンロード", data=csv_bytes, file_name="leads.csv", mime="text/csv")
        with c2:
            try:
                import openpyxl  # for export
                excel_buf = io.BytesIO()
                with pd.ExcelWriter(excel_buf, engine="openpyxl") as writer:
                    leads_df.to_excel(writer, index=False, sheet_name="Leads")
                    assess_df.to_excel(writer, index=False, sheet_name="Assessments")
                st.download_button("Leads/AssessmentsをExcel(.xlsx)でダウンロード", data=excel_buf.getvalue(),
                                   file_name="export.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception as e:
                st.warning(f"Excel書き出しでエラー：{e}（requirements.txt の openpyxl を確認）")

        st.caption("注意：Streamlit Cloudの無料枠では、再デプロイやスリープでローカルDBが初期化されることがあります。永続保存が必要なら外部DB（例：Supabase, Neon等）をご検討ください。")

    # ==== 5) LINEステップ ====
    with tabs[4]:
        st.subheader("見込み客→予約までのLINEステップ配信テンプレ")
        name_placeholder = st.text_input("差し込み用：お名前（任意）", "（お名前）")
        template = build_line_step_template(name_placeholder)
        st.text_area("プレビュー", template, height=300)
        st.download_button("テンプレTXTをダウンロード", data=template.encode("utf-8-sig"),
                           file_name="line_step_template.txt", mime="text/plain")


if __name__ == "__main__":
    main()
